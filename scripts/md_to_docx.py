"""One-shot converter for PHASE_3_SKETCH.md → .docx for offline review."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "PHASE_3_SKETCH.md"
DST = ROOT / "PHASE_3_SKETCH.docx"

doc = Document()

for style_name, size in [("Normal", 11), ("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 13)]:
    s = doc.styles[style_name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)

for section in doc.sections:
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)


def add_inline(paragraph, text):
    """Render bold (**x**) and inline code (`x`) inside a paragraph."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


lines = SRC.read_text(encoding="utf-8").splitlines()

i = 0
while i < len(lines):
    line = lines[i]

    if line.strip() == "":
        i += 1
        continue

    if line.startswith("---"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("— — —").italic = True
        i += 1
        continue

    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=1)
        i += 1
        continue
    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=2)
        i += 1
        continue
    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue

    # Fenced code block
    if line.startswith("```"):
        i += 1
        buf = []
        while i < len(lines) and not lines[i].startswith("```"):
            buf.append(lines[i])
            i += 1
        i += 1  # skip closing fence
        p = doc.add_paragraph()
        run = p.add_run("\n".join(buf))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        continue

    # Table block
    if line.lstrip().startswith("|") and i + 1 < len(lines) and re.match(r"\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
        header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
        i += 2
        rows = []
        while i < len(lines) and lines[i].lstrip().startswith("|"):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
            i += 1
        table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
        table.style = "Light Grid Accent 1"
        for idx, h in enumerate(header_cells):
            cell = table.rows[0].cells[idx]
            cell.text = ""
            add_inline(cell.paragraphs[0], h)
            for r in cell.paragraphs[0].runs:
                r.bold = True
        for ri, row in enumerate(rows, start=1):
            for ci, val in enumerate(row):
                cell = table.rows[ri].cells[ci]
                cell.text = ""
                add_inline(cell.paragraphs[0], val)
        doc.add_paragraph()
        continue

    # Bullet (including nested)
    m = re.match(r"^(\s*)([-*])\s+(.*)$", line)
    if m:
        indent_spaces = len(m.group(1))
        level = min(indent_spaces // 2, 4)
        content = m.group(3)
        # Continuation lines for this bullet
        i += 1
        while i < len(lines) and lines[i].startswith(" " * (indent_spaces + 2)) and not re.match(r"^\s*[-*]\s+", lines[i]):
            content += " " + lines[i].strip()
            i += 1
        style = "List Bullet" if level == 0 else f"List Bullet {level + 1}"
        try:
            p = doc.add_paragraph(style=style)
        except KeyError:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            p.add_run("• ")
        add_inline(p, content)
        continue

    # Numbered list
    m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
    if m:
        indent_spaces = len(m.group(1))
        content = m.group(3)
        i += 1
        while i < len(lines) and lines[i].startswith(" " * (indent_spaces + 3)) and not re.match(r"^\s*(\d+\.|[-*])\s+", lines[i]):
            content += " " + lines[i].strip()
            i += 1
        p = doc.add_paragraph(style="List Number")
        add_inline(p, content)
        continue

    # Plain paragraph — merge following non-blank, non-special lines
    buf = [line]
    i += 1
    while i < len(lines) and lines[i].strip() and not re.match(r"^(#|---|```|\s*[-*]\s|\s*\d+\.\s|\|)", lines[i]):
        buf.append(lines[i])
        i += 1
    p = doc.add_paragraph()
    add_inline(p, " ".join(s.strip() for s in buf))

doc.save(DST)
print(f"Wrote {DST}")
